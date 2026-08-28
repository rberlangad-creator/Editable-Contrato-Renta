import os
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_table_borders(table, color="CCCCCC", sz="4", val="single"):
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>\n'
        f'  <w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>\n'
        f'  <w:left w:val="none"/>\n'
        f'  <w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>\n'
        f'  <w:right w:val="none"/>\n'
        f'  <w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>\n'
        f'  <w:insideV w:val="none"/>\n'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

def build_docx(md_path, docx_path):
    doc = Document()
    
    # Page Setup - Margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Styles
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
    normal_style.paragraph_format.line_spacing = 1.15
    normal_style.paragraph_format.space_after = Pt(6)

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    in_table = False
    table_lines = []

    def flush_table(lines_list):
        if not lines_list:
            return
        # Parse markdown table
        rows_data = []
        for line in lines_list:
            if re.match(r'^\s*\|?\s*:?-+:?\s*\|', line):
                continue
            cells = [c.strip() for c in line.strip().strip('|').split('|')]
            rows_data.append(cells)
            
        if not rows_data:
            return
            
        num_cols = max(len(r) for r in rows_data)
        table = doc.add_table(rows=len(rows_data), cols=num_cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_borders(table, color="D0D7DE", sz="4")
        
        for i, row in enumerate(rows_data):
            tr = table.rows[i]
            is_header = (i == 0)
            
            for j, cell_text in enumerate(row):
                if j < len(tr.cells):
                    cell = tr.cells[j]
                    cell.text = ""
                    p = cell.paragraphs[0]
                    p.paragraph_format.space_after = Pt(2)
                    p.paragraph_format.space_before = Pt(2)
                    p.paragraph_format.line_spacing = 1.05
                    
                    # Bold header
                    run = p.add_run(cell_text.replace('**', ''))
                    if is_header or '**' in cell_text:
                        run.bold = True
                        
                    if is_header:
                        set_cell_background(cell, "1F4E78")
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                        run.font.size = Pt(9.5)
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    else:
                        run.font.size = Pt(9)
                        if i % 2 == 1:
                            set_cell_background(cell, "F9FAFB")
                        else:
                            set_cell_background(cell, "FFFFFF")
                            
                        # Align numbers or short text
                        if re.match(r'^\$?\d', cell_text):
                            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                        elif cell_text in ['Electric', 'Gas LP', '480V', '220V', '110V', '440V', 'Forklift', 'Tugger Train', 'Pallet Jack', 'Pallet Stacker'] or len(cell_text) < 5:
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        else:
                            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    set_cell_margins(cell, top=60, bottom=60, left=100, right=100)

        # Add empty space after table
        p_after = doc.add_paragraph()
        p_after.paragraph_format.space_after = Pt(6)

    i = 0
    while i < len(lines):
        line = lines[i]
        raw_line = line.rstrip('\n')
        stripped = raw_line.strip()

        # Handle Markdown Tables
        if stripped.startswith('|') and stripped.endswith('|'):
            if not in_table:
                in_table = True
                table_lines = []
            table_lines.append(stripped)
            i += 1
            continue
        else:
            if in_table:
                flush_table(table_lines)
                in_table = False
                table_lines = []

        if not stripped:
            i += 1
            continue

        # Horizontal Rule
        if stripped in ['---', '***', '___']:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="6" w:space="1" w:color="D0D7DE"/></w:pBdr>')
            p._p.get_or_add_pPr().append(pBdr)
            i += 1
            continue

        # Headings
        if stripped.startswith('# '):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(12)
            run = p.add_run(stripped[2:].strip().replace('**', ''))
            run.bold = True
            run.font.size = Pt(16)
            run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x78)
            i += 1
            continue
        elif stripped.startswith('## '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(8)
            run = p.add_run(stripped[3:].strip().replace('**', ''))
            run.bold = True
            run.font.size = Pt(13)
            run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
            i += 1
            continue
        elif stripped.startswith('### '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(stripped[4:].strip().replace('**', ''))
            run.bold = True
            run.font.size = Pt(11.5)
            run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x78)
            i += 1
            continue

        # Code Blocks (Signatures block)
        if stripped.startswith('```'):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i].rstrip('\n'))
                i += 1
            i += 1
            
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(8)
            for cl in code_lines:
                run = p.add_run(cl + '\n')
                run.font.name = 'Calibri'
                run.font.size = Pt(10.5)
                run.bold = True
            continue

        # Bullet list items
        if stripped.startswith('- ') or stripped.startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(3)
            content = stripped[2:].strip()
            # process formatting inline
            parts = re.split(r'(\*\*.*?\*\*)', content)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)
            i += 1
            continue

        # Regular Paragraph
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.15

        # Process inline markdown bold
        parts = re.split(r'(\*\*.*?\*\*)', stripped)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = p.add_run(part[2:-2])
                run.bold = True
                if "BLINDAJE" in part or "COBRO CONTINUO" in part or "SEGURO OBLIGATORIO" in part or "PENA CONVENCIONAL" in part:
                    run.font.color.rgb = RGBColor(0xA6, 0x1C, 0x1C)
            else:
                p.add_run(part)

        i += 1

    if in_table:
        flush_table(table_lines)

    doc.save(docx_path)
    print(f"Document saved successfully at: {docx_path}")

if __name__ == "__main__":
    md_file = "/Users/flavioberlanga/.gemini/antigravity/scratch/contrato_milwaukee_ernesto/contrato_arrendamiento_consolidado_milwaukee_ernesto.md"
    docx_file = "/Users/flavioberlanga/.gemini/antigravity/scratch/contrato_milwaukee_ernesto/Contrato_Consolidado_Arrendamiento_Milwaukee_Ernesto.docx"
    build_docx(md_file, docx_file)
